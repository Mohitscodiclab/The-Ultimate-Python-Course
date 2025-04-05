from fastapi import FastAPI, Form, Request
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from docx import Document
from tempfile import NamedTemporaryFile

app = FastAPI()
app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")


@app.get("/")
def index(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})


@app.post("/generate")
async def generate_cv(
    fullName: str = Form(...),
    email: str = Form(...),
    phone: str = Form(...),
    github: str = Form(""),
    portfolio: str = Form(""),
    objective: str = Form(...),
    education: str = Form(...),
    skills: str = Form(...),
    projects: str = Form(...),
    strengths: str = Form(...),
    details: str = Form(...),
    format: str = Form("docx")
):
    # Create DOCX
    doc = Document()
    doc.add_heading(fullName, 0)
    doc.add_paragraph(f"{email} | {phone}")
    doc.add_paragraph(f"GitHub: {github} | Portfolio: {portfolio}")
    doc.add_heading("Career Objective", level=1)
    doc.add_paragraph(objective)
    doc.add_heading("Education", level=1)
    doc.add_paragraph(education)
    doc.add_heading("Technical Skills", level=1)
    doc.add_paragraph(skills)
    doc.add_heading("Projects", level=1)
    doc.add_paragraph(projects)
    doc.add_heading("Strengths", level=1)
    doc.add_paragraph(strengths)
    doc.add_heading("Personal Details", level=1)
    doc.add_paragraph(details)
    doc.add_heading("Declaration", level=1)
    doc.add_paragraph("I hereby declare that the above information is true.\n\nSignature:\n" + fullName)

    with NamedTemporaryFile(delete=False, suffix=f".{format}") as tmp:
        file_path = tmp.name
        if format == "docx":
            doc.save(file_path)
        # PDF and PNG generation can go here later (using pdfkit or Pillow)
        else:
            doc.save(file_path)  # fallback

    return FileResponse(file_path, filename=f"{fullName}_CV.{format}", media_type="application/octet-stream")
